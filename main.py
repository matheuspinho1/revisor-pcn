import os
import datetime
import pdfplumber
from openai import AzureOpenAI
from api_config import API_KEY, ENDPOINT, API_VERSION
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

client = AzureOpenAI(
    azure_endpoint=ENDPOINT,
    api_key=API_KEY,
    api_version=API_VERSION
)

def ler_pdf(caminho, cache_dir="CACHE"):
    """Lê o conteúdo de um arquivo PDF e retorna o texto."""
    # Verificar se já existe no cache 
    nome_arquivo = os.path.basename(caminho)
    if not os.path.exists(cache_dir):
        os.makedirs(cache_dir)
        
    cache_path = os.path.join(cache_dir, f"{nome_arquivo}.txt")
    
    if os.path.exists(cache_path):
        print(f"Usando cache para {nome_arquivo}")
        with open(cache_path, "r", encoding="utf-8") as f:
            return f.read()
    
    print(f"Extraindo texto de {nome_arquivo}...")
    texto = ""
    try:
        with pdfplumber.open(caminho) as pdf:
            total_paginas = len(pdf.pages)
            for i, pagina in enumerate(pdf.pages):
                if i % 20 == 0:  # Mostrar progresso a cada 20 páginas
                    print(f"  Processando página {i+1}/{total_paginas}")
                conteudo = pagina.extract_text()
                if conteudo:
                    texto += conteudo + "\n\n"
        
        # Salvar no cache
        with open(cache_path, "w", encoding="utf-8") as f:
            f.write(texto)
            
        return texto
    except Exception as e:
        print(f"Erro ao processar {nome_arquivo}: {str(e)}")
        return f"ERRO NA EXTRAÇÃO: {str(e)}"

def ler_prompt(caminho):
    """Lê o conteúdo do arquivo de prompt."""
    with open(caminho, "r", encoding="utf-8") as arquivo:
        return arquivo.read()

def chamar_gpt(messages, max_tentativas=3):
    """Chama a API com retry em caso de falha."""
    for tentativa in range(max_tentativas):
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                temperature=0.5 
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"Erro na chamada à API (tentativa {tentativa+1}): {str(e)}")
            if tentativa == max_tentativas - 1:
                raise
            import time
            time.sleep(5)  # Espera 5 segundos antes de tentar novamente

def extrair_estrutura_relatorio(prompt_texto):
    """Extrai a estrutura do relatório a partir do prompt."""
    print("Extraindo estrutura do relatório...")
    
    mensagens = [
        {"role": "system", "content": "Você é um pedagogo especialista em Educação Profissional do setor secundário (Comércio de Bens e Serviços). Sua tarefa é identificar a estrutura proposta de um relatório de análise de Plano de Curso Nacional a partir de um prompt."},
        {"role": "user", "content": f"""
Analise o prompt abaixo e liste APENAS os títulos dos 24 itens que devem compor o relatório final.
Forneça apenas a numeração e o título de cada item, sem explicações adicionais.

PROMPT:
{prompt_texto[:60000]}

Formato esperado:
1. Título do Item 1
2. Título do Item 2
...
24. Título do Item 24
"""}
    ]
    
    try:
        estrutura = chamar_gpt(mensagens)
        print(" Estrutura do relatório extraída com sucesso")
        
        # Verificar se todos os 24 itens estão presentes
        itens = [linha for linha in estrutura.split('\n') if linha.strip()]
        num_itens = len(itens)
        print(f"  → {num_itens} itens identificados na estrutura")
        
        if num_itens < 15:  # Se não conseguiu extrair ao menos 15 itens, algo está errado
            print("  ⚠️ Poucos itens identificados. Tentando abordagem alternativa...")
            # Tente uma abordagem mais direta
            estrutura = """1. Cabeçalho
2. Tabela de Impacto por cada Unidade Curricular (UC)
3. Tabela com propostas de novos nomes para cada UC
4. Perfil Profissional de Conclusão
5. Sugestões de Atualizações para o Perfil Profissional de Conclusão
6. Tabela Comparativa de Carga Horária por cada UC
7. Tabela com alterações dos “Indicadores” de cada competência/Unidade Curricular (UC)
8. Tabela com alterações dos “Conhecimentos” de cada competência/Unidade Curricular (UC)
9. Tabela com alterações das “Habilidades” de cada competência/Unidade Curricular (UC)
10. 10. Tabela com alterações das “Atitudes/Valores” de cada competência/Unidade Curricular (UC)
11. Principal tipo de tecnologias que impactam as competências da profissão
12. Justificativa da resposta do item 11 (Principal tipo de tecnologias que impactam as competências da profissão)
13. Principal tipo de impacto nas competências da profissão
14. Justificativa da resposta do item 13 (Principal tipo de impacto nas competências da profissão)
15. O PCN deve ser Mantido, Atualizado ou Descontinuado?
16. Justificativa da resposta do item 15 (PCN deve ser Mantido, Atualizado ou Descontinuado?).
17. Qual o horizonte de atualização do PCN (caso a resposta do item 13 seja ‘Atualizado’)?
18.Justificativa da resposta do item 176 (horizonte de atualização do PCN).
19. Sugestão do nome de novo curso (caso seja sugerida a descontinuidade do PCN avaliado no item 13)
20. Justificativa da resposta do item 19 (Sugestão do nome de novo curso).
21. Sugestões para atualização de PTDs
22. Projeto Integrador
23. Instalações, Equipamentos e Recursos Didáticos
24. Tabela de referências
"""
        
        return estrutura
    except Exception as e:
        print(f"Erro ao extrair estrutura: {str(e)}")
        # Retornar uma estrutura padrão em caso de falha
        return """1. Cabeçalho
2. Tabela de Impacto por cada Unidade Curricular (UC)
3. Tabela com propostas de novos nomes para cada UC
4. Perfil Profissional de Conclusão
5. Sugestões de Atualizações para o Perfil Profissional de Conclusão
6. Tabela Comparativa de Carga Horária por cada UC
7. Tabela com alterações dos “Indicadores” de cada competência/Unidade Curricular (UC)
8. Tabela com alterações dos “Conhecimentos” de cada competência/Unidade Curricular (UC)
9. Tabela com alterações das “Habilidades” de cada competência/Unidade Curricular (UC)
10. 10. Tabela com alterações das “Atitudes/Valores” de cada competência/Unidade Curricular (UC)
11. Principal tipo de tecnologias que impactam as competências da profissão
12. Justificativa da resposta do item 11 (Principal tipo de tecnologias que impactam as competências da profissão)
13. Principal tipo de impacto nas competências da profissão
14. Justificativa da resposta do item 13 (Principal tipo de impacto nas competências da profissão)
15. O PCN deve ser Mantido, Atualizado ou Descontinuado?
16. Justificativa da resposta do item 15 (PCN deve ser Mantido, Atualizado ou Descontinuado?).
17. Qual o horizonte de atualização do PCN (caso a resposta do item 13 seja ‘Atualizado’)?
18.Justificativa da resposta do item 176 (horizonte de atualização do PCN).
19. Sugestão do nome de novo curso (caso seja sugerida a descontinuidade do PCN avaliado no item 13)
20. Justificativa da resposta do item 19 (Sugestão do nome de novo curso).
21. Sugestões para atualização de PTDs
22. Projeto Integrador
23. Instalações, Equipamentos e Recursos Didáticos
24. Tabela de referências
"""

def extrair_unidades_curriculares(documento_pc):
    """Extrai a lista de Unidades Curriculares do documento."""
    print("Extraindo Unidades Curriculares do documento...")
    
    mensagens = [
        {"role": "system", "content": "Você é um pedagogo especialista em Educação Profissional do setor secundário (Comércio de Bens e Serviços). Identifique precisamente as Unidades Curriculares listadas no documento."},
        {"role": "user", "content": f"""
Analise o documento abaixo e extraia APENAS a lista completa de Unidades Curriculares (UCs) mencionadas.
Para cada UC, forneça o número e o nome/título exato como aparece no documento.

DOCUMENTO:
{documento_pc[:30000]}

Formato esperado:
UC1: [Nome da UC1]
UC2: [Nome da UC2]
...
"""}
    ]
    
    try:
        ucs = chamar_gpt(mensagens)
        print(" Unidades Curriculares extraídas com sucesso")
        return ucs
    except Exception as e:
        print(f"Erro ao extrair UCs: {str(e)}")
        return "Não foi possível extrair as Unidades Curriculares automaticamente."

def gerar_item_relatorio(numero_item, titulo_item, prompt_texto, documentos_base, documento_pc, nome_arquivo_pc, ucs, itens_anteriores=None):
    """Gera um item específico do relatório."""
    print(f"Gerando item {numero_item}: {titulo_item}...")
    
    # Preparar as instruções específicas com base no número do item
    instrucoes_especificas = ""
    contexto_anterior = ""
    
    # Adicionar contexto de itens anteriores quando relevante
    if itens_anteriores and numero_item > 9:
        # Para justificativas, incluir o item que está sendo justificado
        if numero_item in [10, 12, 14, 16, 18]:
            item_referencia = numero_item - 1
            if str(item_referencia) in itens_anteriores:
                contexto_anterior = f"Item {item_referencia} (que você está justificando): {itens_anteriores[str(item_referencia)]}\n\n"
                
    
    # Instruções específicas por tipo de item
    if numero_item in [2, 3, 6, 7, 8, 9, 10, 24]:  # Itens com tabelas
        instrucoes_especificas = """
IMPORTANTE: 
- Formate a tabela corretamente
- Inclua TODAS as Unidades Curriculares do curso na tabela
- Forneça apenas o título e a tabela, sem explicações adicionais
- Na tabela comparativa de carg ahorária por cada uc, crie a tabela mostrando a carga horária anterior e a nova, com o campo para justificativa da nova carga horária
"""
    elif numero_item in [11, 13, 15, 17, 19]:  # Itens de escolha única
        instrucoes_especificas = """
Forneça apenas a resposta direta conforme as opções permitidas, sem explicações adicionais.
"""
    elif numero_item in [12, 14, 16, 18, 20]:  # Itens de justificativa
        instrucoes_especificas = """
Forneça uma justificativa detalhada em texto por extenso com aproximadamente 2.000 caracteres.
"""
    elif numero_item == 1:  # Cabeçalho
        instrucoes_especificas = """
Inclua todos os subitens solicitados: nome do curso, carga horária, eixo tecnológico, segmento profissional, 
ano de criação/revisão, quantidade de UCs e lista completa das UCs.
"""
    elif numero_item == 21:  # Sugestões para PTDs
        instrucoes_especificas = """
Escreva sugestões de orientações para que os docentes atualizem seus Planos de Trabalho Docente (PTDs) até o novo PCN entrar em vigor (estas orientações seriam enviadas para os supervisores pedagógicos). A resposta deve ser em texto por extenso com aproximadamente 2.000 caracteres.
Texto introdutório sobre o profissional e suas atribuições;
Campo de atuação e formas de interação com outros profissionais e setores.
"""

    elif numero_item == 4:  # Perfil Profissional de Conclusão
        instrucoes_especificas = """
      Listar o Perfil Profissional de Conclusão conforme descrito no plano de curso.
"""

    elif numero_item == 5:  # Perfil Profissional de Conclusão Detalhamento
        instrucoes_especificas = """
        - Ao revisar o Plano de Curso, item 4 (Perfil Profissional de Conclusão), adicionar uma seção "Sugestões de Atualizações para o Perfil Profissional de Conclusão", indique atualizações considerando possíveis mudanças no contexto de atuação profissional, inovações do setor ou novos arranjos produtivos e organizacionais(Considere o documento impacto da automação).
As propostas de atualização devem se limitar aos seguintes trechos:
Texto introdutório sobre o profissional e suas atribuições;
Campo de atuação e formas de interação com outros profissionais e setores.
"""

    elif numero_item == 23:  # Instalações, Equipamentos e Recursos Didáticos
        instrucoes_especificas = """
        Atualizações nos Instalações, Equipamentos e Recursos Didáticos: 
Inserir as demandas por equipamentos e recursos necessários à implementação das atualizações sugeridas nas Unidades Curriculares. Caso alguma UC seja impactada por transformações decorrentes da automação, devem ser especificados os recursos e adequações requeridos nos ambientes de aprendizagem para sua viabilização. 
        """
    
    # Adicionar informações sobre UCs quando relevante
    if numero_item in [2, 3, 6, 7, 8, 9, 10]:
        instrucoes_especificas += f"\n\nUNIDADES CURRICULARES DO CURSO:\n{ucs}\n"
    
    mensagens = [
        {"role": "system", "content": f"Você é um especialista em educação profissional do Senac. Gere APENAS o item {numero_item} do relatório conforme solicitado."},
        {"role": "user", "content": f"""
Gere APENAS o item {numero_item} ({titulo_item}) do relatório para o PCN "{nome_arquivo_pc}", seguindo exatamente o formato solicitado.

{contexto_anterior}
{instrucoes_especificas}

ESTRUTURA DO RELATÓRIO:
Item {numero_item}: {titulo_item}

DOCUMENTO PC:
{documento_pc[:30000]}

DOCUMENTOS DE REFERÊNCIA:
{next(iter(documentos_base.values()))[:100000]}

PROMPT ORIGINAL:
{prompt_texto[:500000]}

Forneça apenas o conteúdo do item {numero_item}, começando com o título "## {numero_item}. {titulo_item}".
"""}
    ]
    
    try:
        conteudo_item = chamar_gpt(mensagens)
        # Verificar se o título está presente e adicionar se necessário
        if not conteudo_item.startswith(f"## {numero_item}.") and not conteudo_item.startswith(f"##{numero_item}."):
            conteudo_item = f"## {numero_item}. {titulo_item}\n\n{conteudo_item}"
        
        print(f" Item {numero_item} gerado com sucesso")
        return conteudo_item
    except Exception as e:
        print(f"Erro ao gerar item {numero_item}: {str(e)}")
        return f"## {numero_item}. {titulo_item}\n\nNão foi possível gerar este item automaticamente."

def gerar_relatorio_completo(prompt_texto, documentos_base, documento_pc, nome_arquivo_pc):
    """Gera o relatório completo com todos os 24 itens garantidos."""
    print("Iniciando geração do relatório completo por itens...")
    
    # 1. Extrair a estrutura do relatório
    estrutura_texto = extrair_estrutura_relatorio(prompt_texto)
    linhas_estrutura = [linha.strip() for linha in estrutura_texto.split('\n') if linha.strip()]
    estrutura = {}
    
    for linha in linhas_estrutura:
        partes = linha.split('. ', 1)
        if len(partes) == 2 and partes[0].isdigit():
            estrutura[partes[0]] = partes[1]
    
    # 2. Extrair as Unidades Curriculares
    ucs = extrair_unidades_curriculares(documento_pc)
    
    # 3. Gerar cada item individualmente
    itens_gerados = {}
    
    # Começar com o título principal
    relatorio_final = f"# Revisão do PCN \"{nome_arquivo_pc}\"\n\n"
    
    # Gerar itens de 1 a 24
    for i in range(1, 25):
        numero = str(i)
        if numero in estrutura:
            titulo = estrutura[numero]
            
            # Gerar o item
            conteudo = gerar_item_relatorio(
                i, 
                titulo, 
                prompt_texto, 
                documentos_base, 
                documento_pc, 
                nome_arquivo_pc,
                ucs,
                itens_gerados
            )
            
            # Salvar o conteúdo para referência e para items de justificativa
            itens_gerados[numero] = conteudo
            
            # Adicionar ao relatório final
            relatorio_final += conteudo + "\n\n"
        else:
            print(f"⚠️ Item {i} não encontrado na estrutura")
            relatorio_final += f"## {i}. Item não definido\n\nEste item não foi encontrado na estrutura do relatório.\n\n"
            
    # 5. Salvar o relatório em texto para referência
    os.makedirs("CACHE", exist_ok=True)
    with open(f"CACHE/relatorio_completo_{nome_arquivo_pc}.md", "w", encoding="utf-8") as f:
        f.write(relatorio_final)
    
    print(" Relatório completo com 24 itens gerado com sucesso")
    return relatorio_final

def processar_tabela_markdown(texto_tabela):
    """Processa uma tabela em formato markdown e retorna linhas e colunas."""
    linhas = texto_tabela.strip().split('\n')
    
    if not linhas or len(linhas) < 3:  # Precisa ter pelo menos cabeçalho, separador e uma linha de dados
        return [], []
        
    # Encontrar a linha de cabeçalho
    header_line = None
    for i, linha in enumerate(linhas):
        if '|' in linha and i < len(linhas) - 1 and '---' in linhas[i+1] and '|' in linhas[i+1]:
            header_line = linha
            break
    
    if not header_line:
        return [], []
    
    # Extrair colunas do cabeçalho
    colunas = []
    for col in header_line.split('|'):
        col = col.strip()
        if col:
            colunas.append(col)
    
    # Pular linha de separação e extrair dados
    dados = []
    data_start = False
    for linha in linhas:
        if '---' in linha and '|' in linha:
            data_start = True
            continue
        
        if data_start and '|' in linha:
            row_data = []
            for col in linha.split('|'):
                col = col.strip()
                row_data.append(col)
            
            # Remover células vazias no início e fim (resultantes da separação por |)
            row_data = [col for col in row_data if col]
            
            if row_data and len(row_data) > 0:
                # Ajustar tamanho da linha para corresponder ao número de colunas
                while len(row_data) < len(colunas):
                    row_data.append("")
                dados.append(row_data[:len(colunas)])
    
    return colunas, dados

def gerar_relatorio_docx(nome_arquivo, conteudo, output_dir="RELATORIOS"):
    """Gera um relatório formatado em DOCX com tabelas e formatação específica."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    nome_saida = f"{output_dir}/{timestamp}_{nome_arquivo}.docx"
    
    doc = Document()
    
    # Configuração básica do documento
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
    
    # Processar o conteúdo linha por linha
    linhas = conteudo.split('\n')
    i = 0
    
    while i < len(linhas):
        linha = linhas[i].strip()
        
        # Título principal (# Revisão do PCN...)
        if linha.startswith('# '):
            heading = doc.add_heading(linha[2:], level=0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
        
        # Cabeçalhos nível 2 (## 1. Cabeçalho)
        elif linha.startswith('## '):
            # Extrair o número e o texto
            texto_heading = linha[3:]
            doc.add_heading(texto_heading, level=1)
            i += 1
        
        # Subcabeçalhos nível 3 (### - Nome do curso)
        elif linha.startswith('### '):
            texto_heading = linha[4:]
            doc.add_heading(texto_heading, level=2)
            i += 1
        
        # Detectar início de tabela
        elif '|' in linha and i+1 < len(linhas) and '|' in linhas[i+1] and ('---' in linhas[i+1] or '—--' in linhas[i+1] or '-—-' in linhas[i+1]):
            # Coletar todas as linhas da tabela
            tabela_markdown = linha + '\n'
            j = i + 1
            while j < len(linhas) and '|' in linhas[j]:
                tabela_markdown += linhas[j] + '\n'
                j += 1
            
            # Processar a tabela markdown
            colunas, dados = processar_tabela_markdown(tabela_markdown)
            
            if colunas and dados:
                # Criar tabela no Word
                table = doc.add_table(rows=1, cols=len(colunas))
                table.style = 'Table Grid'
                
                # Adicionar cabeçalhos
                head_cells = table.rows[0].cells
                for idx, col in enumerate(colunas):
                    head_cells[idx].text = col
                    # Aplicar negrito ao texto do cabeçalho
                    for paragraph in head_cells[idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Adicionar linhas de dados
                for row_data in dados:
                    cells = table.add_row().cells
                    for idx, text in enumerate(row_data):
                        if idx < len(cells):
                            cells[idx].text = text
                
                # Aplicar formatação na tabela - abordagem simplificada que evita o bug
                table.style = 'Table Grid'  # Esta linha aplica bordas simples
                
                # Ajustar largura da tabela para preencher a página
                table.autofit = True
                
                # Adicionar espaço após a tabela
                doc.add_paragraph()
                
                # Avançar para após a tabela
                i = j
            else:
                # Se falhou em processar a tabela, tratar como texto normal
                p = doc.add_paragraph()
                p.text = linha
                i += 1
        
        # Listas com marcadores
        elif linha.startswith('- '):
            p = doc.add_paragraph()
            p.style = 'List Bullet'
            p.text = linha[2:]  # Remover o marcador '- '
            i += 1
        
        # Texto normal
        elif linha:
            p = doc.add_paragraph()
            p.text = linha
            i += 1
        
        # Linha vazia
        else:
            i += 1
    
    # Adicionar rodapé
    last_paragraph = doc.add_paragraph()
    last_paragraph.text = "DESENVOLVIDO POR SENAC DEPARTAMENTO NACIONAL - GER. DE TECNOLOGIAS E DESENHOS EDUCACIONAIS."
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in last_paragraph.runs:
        run.bold = True
    
    # Salvar o documento
    doc.save(nome_saida)
    print(f"Relatório DOCX gerado: {nome_saida}")
    return nome_saida

def main():
    """Função principal do programa."""
    print("\n" + "="*50)
    print("INICIANDO PROCESSAMENTO")
    print("="*50)
    
    # Configurar diretórios
    base_dir = "BASE"
    pc_dir = "PC"
    prompt_file = "prompt.txt"
    
    # 1. Ler prompt uma única vez
    print("\nLendo o prompt...")
    prompt_texto = ler_prompt(prompt_file)
    print(f"Prompt lido: {len(prompt_texto)} caracteres")
    
    # 2. Ler documentos base uma única vez
    print("\nLendo documentos base...")
    documentos_base = {}
    for arquivo in os.listdir(base_dir):
        if arquivo.endswith(".pdf"):
            caminho = os.path.join(base_dir, arquivo)
            documentos_base[arquivo] = ler_pdf(caminho)
            print(f"Documento base lido: {arquivo} ({len(documentos_base[arquivo])} caracteres)")
    
    # 3. Processar cada documento PC
    print("\nProcessando documentos PC...")
    for arquivo in os.listdir(pc_dir):
        if arquivo.endswith(".pdf"):
            try:
                print(f"\nProcessando: {arquivo}")
                caminho = os.path.join(pc_dir, arquivo)
                
                # Ler o documento PC uma única vez
                documento_pc = ler_pdf(caminho)
                print(f" Documento PC lido: {len(documento_pc)} caracteres")
                
                # Gerar o relatório completo com todos os 23 itens garantidos
                print("\nGerando relatório completo com 23 itens...")
                relatorio_texto = gerar_relatorio_completo(
                    prompt_texto, 
                    documentos_base, 
                    documento_pc,
                    arquivo.replace(".pdf", "")
                )
                
                # Gerar documento Word formatado
                print("\nGerando documento Word formatado...")
                nome_saida = gerar_relatorio_docx(arquivo.replace(".pdf", ""), relatorio_texto)
                
                print(f"\n Processamento completo: {arquivo} → {nome_saida}")
            except Exception as e:
                print(f"\n❌ ERRO ao processar {arquivo}: {str(e)}")
                import traceback
                traceback.print_exc()
    
    print("\n" + "="*50)
    print("PROCESSAMENTO CONCLUÍDO")
    print("="*50)

if __name__ == "__main__":
    main()